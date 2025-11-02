"""
Advanced Document Parser with Metadata Extraction
Extracts structure, sections, entities, and metadata from legal documents
"""
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

import pdfplumber
from docx import Document as DocxDocument
import dateparser

from backend.models.knowledge_schema import (
    StructuredDocument,
    EnrichedChunk,
    SectionNode,
    DocType,
    Jurisdiction,
    DocumentVersion
)
from backend.utils.logger import setup_logger

logger = setup_logger(__name__)


class LegalDocumentParser:
    """Advanced parser for legal documents"""

    # Section patterns (common in legal docs)
    SECTION_PATTERNS = [
        r'^(?:ARTICLE|Article)\s+([IVXLCDM]+|\d+)[:\.\s]+(.+)$',  # ARTICLE I: Title
        r'^(?:§|Section|SECTION)\s*(\d+(?:\.\d+)*)[:\.\s]+(.+)$',  # §5.2: Title
        r'^(\d+(?:\.\d+)*)[:\.\s]+([A-Z][^\.]+)$',  # 5.2: Title
        r'^([A-Z\s]{3,}):?\s*$',  # ALL CAPS HEADERS
    ]

    # Date patterns
    DATE_PATTERNS = [
        r'(?:effective|executed|signed|dated)(?:\s+as\s+of)?\s*:?\s*([A-Z][a-z]+\s+\d{1,2},?\s+\d{4})',
        r'(?:effective|executed|signed|dated)(?:\s+as\s+of)?\s*:?\s*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})',
    ]

    # Party patterns
    PARTY_PATTERNS = [
        r'between\s+([A-Z][^,\(]+?)(?:\s+\([^\)]+\))?\s+and\s+([A-Z][^,\(]+?)(?:\s+\([^\)]+\))?',
        r'by\s+and\s+between\s+([A-Z][^,\(]+?)(?:\s+\([^\)]+\))?\s+and\s+([A-Z][^,\(]+?)(?:\s+\([^\)]+\))?',
    ]

    # Jurisdiction patterns
    JURISDICTION_PATTERNS = {
        'CA': r'\b(?:California|State\s+of\s+California)\b',
        'NY': r'\b(?:New\s+York|State\s+of\s+New\s+York)\b',
        'TX': r'\b(?:Texas|State\s+of\s+Texas)\b',
        'FL': r'\b(?:Florida|State\s+of\s+Florida)\b',
        'US': r'\b(?:United\s+States|Federal|U\.S\.)\b',
    }

    # Document type patterns
    DOCTYPE_PATTERNS = {
        DocType.CONTRACT: r'\b(?:contract|agreement|employment\s+agreement)\b',
        DocType.NDA: r'\b(?:non-disclosure|confidentiality\s+agreement|NDA)\b',
        DocType.POLICY: r'\b(?:policy|handbook|procedures?)\b',
        DocType.LICENSE: r'\b(?:license|licensing\s+agreement)\b',
        DocType.MEMO: r'\b(?:memorandum|memo)\b',
    }

    def __init__(self):
        pass

    def parse_document(self, file_path: Path) -> Optional[StructuredDocument]:
        """
        Parse document based on file extension

        Args:
            file_path: Path to document file

        Returns:
            StructuredDocument or None if parsing fails
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        extension = file_path.suffix.lower()

        try:
            if extension == '.pdf':
                return self.parse_pdf(file_path)
            elif extension in ['.txt', '.md']:
                return self.parse_txt(file_path)
            elif extension in ['.docx', '.doc']:
                return self.parse_docx(file_path)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return None
        except Exception as e:
            logger.error(f"Error parsing {file_path.name}: {e}")
            return None

    def parse_pdf(self, file_path: Path) -> StructuredDocument:
        """Parse PDF with advanced structure extraction"""
        logger.info(f"Advanced parsing: {file_path.name}")

        # Calculate file hash
        file_hash = self._calculate_hash(file_path)

        # Extract text with structure
        full_text = ""
        pages_text = []

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages_text.append(text)
                full_text += text + "\n\n"

        # Extract metadata
        title = self._extract_title(full_text, file_path.stem)
        doctype, doctype_conf = self._classify_document(full_text)
        jurisdiction, juris_conf = self._extract_jurisdiction(full_text)
        parties = self._extract_parties(full_text)
        dates = self._extract_dates(full_text)
        defined_terms = self._extract_definitions(full_text)

        # Extract sections
        sections, section_tree = self._extract_sections(pages_text)

        # Create chunks with context
        chunks = self._create_enriched_chunks(
            pages_text,
            sections,
            file_path.stem,
            file_hash
        )

        # Detect key clauses
        key_clauses = self._detect_key_clauses(full_text)

        # Create structured document
        doc = StructuredDocument(
            doc_id=file_hash[:12],
            title=title,
            file_path=str(file_path),
            file_hash=file_hash,
            doctype=doctype,
            doctype_confidence=doctype_conf,
            jurisdiction=jurisdiction,
            jurisdiction_confidence=juris_conf,
            parties=parties,
            creation_date=dates.get('creation'),
            effective_date=dates.get('effective'),
            expiration_date=dates.get('expiration'),
            indexed_at=datetime.utcnow(),
            version=DocumentVersion.SIGNED,
            total_pages=total_pages,
            total_sections=len(sections),
            section_tree=section_tree,
            chunks=chunks,
            total_chunks=len(chunks),
            full_text=full_text,
            defined_terms=defined_terms,
            key_clauses=key_clauses,
        )

        logger.info(
            f"Parsed: {title} | Type: {doctype} | "
            f"Jurisdiction: {jurisdiction} | Sections: {len(sections)} | Chunks: {len(chunks)}"
        )

        return doc

    def parse_txt(self, file_path: Path) -> StructuredDocument:
        """Parse TXT file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            full_text = f.read()

        file_hash = self._calculate_hash(file_path)

        # Split into pages (simulate)
        pages_text = self._split_text_into_pages(full_text)

        # Extract metadata
        title = self._extract_title(full_text, file_path.stem)
        doctype, doctype_conf = self._classify_document(full_text)
        jurisdiction, juris_conf = self._extract_jurisdiction(full_text)
        parties = self._extract_parties(full_text)
        dates = self._extract_dates(full_text)
        defined_terms = self._extract_definitions(full_text)

        # Extract sections
        sections, section_tree = self._extract_sections(pages_text)

        # Create chunks
        chunks = self._create_enriched_chunks(
            pages_text,
            sections,
            file_path.stem,
            file_hash
        )

        key_clauses = self._detect_key_clauses(full_text)

        doc = StructuredDocument(
            doc_id=file_hash[:12],
            title=title,
            file_path=str(file_path),
            file_hash=file_hash,
            doctype=doctype,
            doctype_confidence=doctype_conf,
            jurisdiction=jurisdiction,
            jurisdiction_confidence=juris_conf,
            parties=parties,
            creation_date=dates.get('creation'),
            effective_date=dates.get('effective'),
            indexed_at=datetime.utcnow(),
            version=DocumentVersion.DRAFT,
            total_pages=len(pages_text),
            total_sections=len(sections),
            section_tree=section_tree,
            chunks=chunks,
            total_chunks=len(chunks),
            full_text=full_text,
            defined_terms=defined_terms,
            key_clauses=key_clauses,
        )

        return doc

    def parse_docx(self, file_path: Path) -> StructuredDocument:
        """Parse DOCX file"""
        try:
            docx_doc = DocxDocument(file_path)

            # Extract full text
            full_text = '\n'.join([para.text for para in docx_doc.paragraphs])

            file_hash = self._calculate_hash(file_path)

            # Split into pages (simulate)
            pages_text = self._split_text_into_pages(full_text)

            # Extract metadata
            title = self._extract_title(full_text, file_path.stem)
            doctype, doctype_conf = self._classify_document(full_text)
            jurisdiction, juris_conf = self._extract_jurisdiction(full_text)
            parties = self._extract_parties(full_text)
            dates = self._extract_dates(full_text)
            defined_terms = self._extract_definitions(full_text)

            # Extract sections
            sections, section_tree = self._extract_sections(pages_text)

            # Create chunks
            chunks = self._create_enriched_chunks(
                pages_text,
                sections,
                file_path.stem,
                file_hash
            )

            key_clauses = self._detect_key_clauses(full_text)

            doc = StructuredDocument(
                doc_id=file_hash[:12],
                title=title,
                file_path=str(file_path),
                file_hash=file_hash,
                doctype=doctype,
                doctype_confidence=doctype_conf,
                jurisdiction=jurisdiction,
                jurisdiction_confidence=juris_conf,
                parties=parties,
                creation_date=dates.get('creation'),
                effective_date=dates.get('effective'),
                indexed_at=datetime.utcnow(),
                version=DocumentVersion.DRAFT,
                total_pages=len(pages_text),
                total_sections=len(sections),
                section_tree=section_tree,
                chunks=chunks,
                total_chunks=len(chunks),
                full_text=full_text,
                defined_terms=defined_terms,
                key_clauses=key_clauses,
            )

            return doc

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path.name}: {e}")
            # Fallback: treat as text
            return self.parse_txt(file_path)

    def _calculate_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()

    def _extract_title(self, text: str, filename: str) -> str:
        """Extract document title"""
        # Look for title in first 500 chars
        lines = text[:500].split('\n')
        for line in lines:
            line = line.strip()
            if len(line) > 10 and len(line) < 100 and line.isupper():
                return line.title()
            elif len(line) > 10 and len(line) < 100:
                # Check if looks like a title
                if any(keyword in line.lower() for keyword in ['agreement', 'contract', 'policy', 'nda']):
                    return line

        return filename.replace('_', ' ').title()

    def _classify_document(self, text: str) -> Tuple[DocType, float]:
        """Classify document type"""
        text_lower = text[:2000].lower()  # First 2000 chars

        scores = {}
        for doctype, pattern in self.DOCTYPE_PATTERNS.items():
            matches = len(re.findall(pattern, text_lower, re.IGNORECASE))
            scores[doctype] = matches

        if not scores or max(scores.values()) == 0:
            return DocType.OTHER, 0.0

        best_type = max(scores, key=scores.get)
        confidence = min(scores[best_type] / 5.0, 1.0)  # Normalize

        return best_type, confidence

    def _extract_jurisdiction(self, text: str) -> Tuple[Jurisdiction, float]:
        """Extract jurisdiction"""
        text_sample = text[:3000]  # First 3000 chars

        scores = {}
        for juris, pattern in self.JURISDICTION_PATTERNS.items():
            matches = len(re.findall(pattern, text_sample, re.IGNORECASE))
            scores[juris] = matches

        if not scores or max(scores.values()) == 0:
            return Jurisdiction.OTHER, 0.0

        best_juris = max(scores, key=scores.get)
        confidence = min(scores[best_juris] / 3.0, 1.0)

        return Jurisdiction(best_juris), confidence

    def _extract_parties(self, text: str) -> List[str]:
        """Extract party names"""
        parties = []
        text_sample = text[:2000]

        for pattern in self.PARTY_PATTERNS:
            matches = re.finditer(pattern, text_sample, re.IGNORECASE)
            for match in matches:
                for group in match.groups():
                    if group:
                        party = group.strip().strip(',').strip()
                        if len(party) > 3 and party not in parties:
                            parties.append(party)

        return parties[:5]  # Max 5 parties

    def _extract_dates(self, text: str) -> Dict[str, Optional[datetime]]:
        """Extract important dates"""
        dates = {
            'creation': None,
            'effective': None,
            'expiration': None
        }

        text_sample = text[:3000]

        for pattern in self.DATE_PATTERNS:
            matches = re.finditer(pattern, text_sample, re.IGNORECASE)
            for match in matches:
                date_str = match.group(1)
                parsed_date = dateparser.parse(date_str)
                if parsed_date:
                    # Heuristic: first date is usually effective date
                    if not dates['effective']:
                        dates['effective'] = parsed_date
                    break

        return dates

    def _extract_definitions(self, text: str) -> Dict[str, str]:
        """Extract defined terms"""
        defined_terms = {}

        # Pattern: "Term" means ...
        pattern = r'"([^"]+)"\s+(?:means|shall\s+mean|is\s+defined\s+as)\s+([^\.]+)\.'
        matches = re.finditer(pattern, text[:5000], re.IGNORECASE)

        for match in matches:
            term = match.group(1).strip()
            definition = match.group(2).strip()
            if len(term) < 50 and len(definition) < 200:
                defined_terms[term] = definition

        return defined_terms

    def _detect_key_clauses(self, text: str) -> List[str]:
        """Detect key clause types"""
        clauses = []
        text_lower = text.lower()

        clause_keywords = {
            'non-compete': r'non-compete|non\s+competition',
            'confidentiality': r'confidential|non-disclosure',
            'termination': r'termination|cancellation',
            'arbitration': r'arbitration|dispute\s+resolution',
            'liability': r'liability|indemnif',
            'intellectual-property': r'intellectual\s+property|ip\s+rights',
            'governing-law': r'governing\s+law|choice\s+of\s+law',
        }

        for clause_type, pattern in clause_keywords.items():
            if re.search(pattern, text_lower):
                clauses.append(clause_type)

        return clauses

    def _extract_sections(self, pages_text: List[str]) -> Tuple[List[Dict], List[SectionNode]]:
        """Extract section structure"""
        sections = []
        section_nodes = []
        current_page = 0
        section_id_counter = 0

        for page_num, page_text in enumerate(pages_text, start=1):
            lines = page_text.split('\n')

            for line_idx, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue

                # Try to match section patterns
                for pattern in self.SECTION_PATTERNS:
                    match = re.match(pattern, line)
                    if match:
                        section_id = f"sec_{section_id_counter}"
                        section_id_counter += 1

                        # Extract number and title
                        if len(match.groups()) >= 2:
                            number = match.group(1).strip()
                            title = match.group(2).strip()
                        else:
                            number = None
                            title = line

                        # Determine level based on number format
                        level = self._determine_section_level(number)

                        section_info = {
                            'id': section_id,
                            'number': number,
                            'title': title,
                            'page': page_num,
                            'level': level
                        }

                        sections.append(section_info)

                        node = SectionNode(
                            id=section_id,
                            number=number,
                            title=title,
                            level=level,
                            page_start=page_num
                        )
                        section_nodes.append(node)

                        break

        # Build hierarchy
        self._build_section_hierarchy(section_nodes)

        return sections, section_nodes

    def _determine_section_level(self, number: Optional[str]) -> int:
        """Determine section level from number format"""
        if not number:
            return 0

        # Count dots to determine depth
        # 1 -> level 1
        # 1.1 -> level 2
        # 1.1.1 -> level 3
        if '.' in number:
            return number.count('.') + 1
        else:
            return 1

    def _build_section_hierarchy(self, nodes: List[SectionNode]):
        """Build parent-child relationships"""
        for i, node in enumerate(nodes):
            # Find parent (previous node with lower level)
            for j in range(i - 1, -1, -1):
                if nodes[j].level < node.level:
                    node.parent_id = nodes[j].id
                    nodes[j].children_ids.append(node.id)
                    break

    def _create_enriched_chunks(
        self,
        pages_text: List[str],
        sections: List[Dict],
        doc_name: str,
        doc_id: str
    ) -> List[EnrichedChunk]:
        """Create enriched chunks with metadata"""
        chunks = []
        chunk_counter = 0

        for page_num, page_text in enumerate(pages_text, start=1):
            # Find current section for this page
            current_section = None
            for section in sections:
                if section['page'] == page_num:
                    current_section = section

            # Split page into chunks (512 chars with 50 overlap)
            page_chunks = self._chunk_text(page_text, size=512, overlap=50)

            for chunk_text in page_chunks:
                if not chunk_text.strip():
                    continue

                chunk_id = f"{doc_id[:12]}#p{page_num}#c{chunk_counter:02d}"
                chunk_counter += 1

                # Analyze chunk
                is_header = self._is_header(chunk_text)
                contains_dates = bool(re.search(r'\d{1,2}[-/]\d{1,2}[-/]\d{2,4}', chunk_text))
                contains_money = bool(re.search(r'\$[\d,]+', chunk_text))

                chunk = EnrichedChunk(
                    chunk_id=chunk_id,
                    doc_id=doc_id[:12],
                    text=chunk_text,
                    tokens=len(chunk_text.split()),
                    page_start=page_num,
                    page_end=page_num,
                    section_id=current_section['id'] if current_section else None,
                    section_title=current_section['title'] if current_section else None,
                    section_path=[current_section['title']] if current_section else [],
                    is_header=is_header,
                    contains_dates=contains_dates,
                    contains_money=contains_money,
                    meta={'doc_name': doc_name}
                )

                chunks.append(chunk)

        return chunks

    def _chunk_text(self, text: str, size: int = 512, overlap: int = 50) -> List[str]:
        """Split text into chunks"""
        if not text.strip():
            return []

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = start + size
            chunk = text[start:end]

            # Try to break at sentence boundary
            if end < text_length:
                last_period = chunk.rfind('. ')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)

                if break_point > size * 0.5:
                    chunk = chunk[:break_point + 1]
                    end = start + break_point + 1

            chunks.append(chunk.strip())
            start = end - overlap

        return [c for c in chunks if c]

    def _is_header(self, text: str) -> bool:
        """Check if text looks like a header"""
        text = text.strip()

        # Short and all caps
        if len(text) < 100 and text.isupper():
            return True

        # Matches section pattern
        for pattern in self.SECTION_PATTERNS:
            if re.match(pattern, text):
                return True

        return False

    def _split_text_into_pages(self, text: str, chars_per_page: int = 3000) -> List[str]:
        """Split text into simulated pages"""
        pages = []
        lines = text.split('\n')
        current_page = []
        current_length = 0

        for line in lines:
            current_page.append(line)
            current_length += len(line)

            if current_length >= chars_per_page:
                pages.append('\n'.join(current_page))
                current_page = []
                current_length = 0

        if current_page:
            pages.append('\n'.join(current_page))

        return pages
